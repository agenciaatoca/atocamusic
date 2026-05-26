from docx import Document
from docx.shared import Pt, Cm, RGBColor, Inches, Emu
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_ALIGN_VERTICAL
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from docx.util import Pt, Cm
import copy

RED   = RGBColor(0xC0, 0x39, 0x2B)
DARK  = RGBColor(0x1A, 0x1A, 0x1A)
MID   = RGBColor(0x44, 0x44, 0x44)
LIGHT = RGBColor(0x88, 0x88, 0x88)
GOLD  = RGBColor(0xC2, 0x7B, 0x00)
WHITE = RGBColor(0xFF, 0xFF, 0xFF)
GBG   = RGBColor(0xF7, 0xF7, 0xF7)
DARK2 = RGBColor(0x2C, 0x2C, 0x2C)

def hex_color(r, g, b):
    return f"{r:02X}{g:02X}{b:02X}"

def set_cell_bg(cell, r, g, b):
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:val'), 'clear')
    shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'), hex_color(r, g, b))
    tcPr.append(shd)

def set_cell_borders(cell, top=None, bottom=None, left=None, right=None):
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    tcBorders = OxmlElement('w:tcBorders')
    for side, val in [('top', top), ('bottom', bottom), ('left', left), ('right', right)]:
        if val is not None:
            el = OxmlElement(f'w:{side}')
            el.set(qn('w:val'), val.get('val', 'single'))
            el.set(qn('w:sz'), str(val.get('sz', 4)))
            el.set(qn('w:color'), val.get('color', '000000'))
            tcBorders.append(el)
        else:
            el = OxmlElement(f'w:{side}')
            el.set(qn('w:val'), 'none')
            el.set(qn('w:sz'), '0')
            el.set(qn('w:color'), 'FFFFFF')
            tcBorders.append(el)
    tcPr.append(tcBorders)

def no_borders(cell):
    set_cell_borders(cell, top={}, bottom={}, left={}, right={})
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    tcBorders = tcPr.find(qn('w:tcBorders'))
    if tcBorders is not None:
        tcPr.remove(tcBorders)
    tcBorders = OxmlElement('w:tcBorders')
    for side in ['top','bottom','left','right']:
        el = OxmlElement(f'w:{side}')
        el.set(qn('w:val'), 'none')
        el.set(qn('w:sz'), '0')
        el.set(qn('w:space'), '0')
        el.set(qn('w:color'), 'FFFFFF')
        tcBorders.append(el)
    tcPr.append(tcBorders)

def set_table_no_borders(table):
    tbl = table._tbl
    tblPr = tbl.tblPr
    tblBorders = OxmlElement('w:tblBorders')
    for side in ['top','left','bottom','right','insideH','insideV']:
        el = OxmlElement(f'w:{side}')
        el.set(qn('w:val'), 'none')
        el.set(qn('w:sz'), '0')
        el.set(qn('w:color'), 'FFFFFF')
        tblBorders.append(el)
    tblPr.append(tblBorders)

def set_cell_margin(cell, top=80, bottom=80, left=120, right=120):
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    tcMar = OxmlElement('w:tcMar')
    for side, val in [('top', top), ('bottom', bottom), ('left', left), ('right', right)]:
        el = OxmlElement(f'w:{side}')
        el.set(qn('w:w'), str(val))
        el.set(qn('w:type'), 'dxa')
        tcMar.append(el)
    tcPr.append(tcMar)

def add_paragraph(doc, text='', bold=False, size=11, color=None, align=WD_ALIGN_PARAGRAPH.LEFT,
                  space_before=6, space_after=6, italic=False, all_caps=False):
    p = doc.add_paragraph()
    p.alignment = align
    p.paragraph_format.space_before = Pt(space_before)
    p.paragraph_format.space_after = Pt(space_after)
    if text:
        run = p.add_run(text)
        run.bold = bold
        run.italic = italic
        run.font.size = Pt(size)
        if color:
            run.font.color.rgb = color
        if all_caps:
            run.font.all_caps = True
    return p

def add_run(paragraph, text, bold=False, size=11, color=None, italic=False):
    run = paragraph.add_run(text)
    run.bold = bold
    run.italic = italic
    run.font.size = Pt(size)
    if color:
        run.font.color.rgb = color
    return run

def add_hr(doc, color_hex='DDDDDD', thickness=4):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(4)
    p.paragraph_format.space_after = Pt(4)
    pPr = p._p.get_or_add_pPr()
    pBdr = OxmlElement('w:pBdr')
    bottom = OxmlElement('w:bottom')
    bottom.set(qn('w:val'), 'single')
    bottom.set(qn('w:sz'), str(thickness))
    bottom.set(qn('w:color'), color_hex)
    bottom.set(qn('w:space'), '1')
    pBdr.append(bottom)
    pPr.append(pBdr)
    return p

def add_red_hr(doc):
    return add_hr(doc, 'C0392B', 10)

def add_section_label(doc, text):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(14)
    p.paragraph_format.space_after = Pt(4)
    run = p.add_run(text)
    run.bold = True
    run.font.size = Pt(8.5)
    run.font.color.rgb = RED
    run.font.all_caps = True
    # letter spacing via XML
    rPr = run._r.get_or_add_rPr()
    spacing = OxmlElement('w:spacing')
    spacing.set(qn('w:val'), '60')
    rPr.append(spacing)
    return p

def set_row_height(row, height_cm):
    tr = row._tr
    trPr = tr.get_or_add_trPr()
    trHeight = OxmlElement('w:trHeight')
    trHeight.set(qn('w:val'), str(int(height_cm * 567)))
    trPr.append(trHeight)

# ─────────────────────────────────────────────────────────────────────────────
doc = Document()

# Page margins
section = doc.sections[0]
section.page_width  = Cm(21)
section.page_height = Cm(29.7)
section.left_margin   = Cm(2.0)
section.right_margin  = Cm(2.0)
section.top_margin    = Cm(1.5)
section.bottom_margin = Cm(1.8)

# Default font
style = doc.styles['Normal']
style.font.name = 'Arial'
style.font.size = Pt(11)

# ── LOGO ──────────────────────────────────────────────────────────────────────
p_logo = doc.add_paragraph()
p_logo.alignment = WD_ALIGN_PARAGRAPH.LEFT
p_logo.paragraph_format.space_before = Pt(0)
p_logo.paragraph_format.space_after  = Pt(0)
run_logo = p_logo.add_run()
run_logo.add_picture('/home/claude/logo_prop2.png', width=Cm(7))

add_paragraph(doc, space_before=2, space_after=2)
add_red_hr(doc)
add_paragraph(doc, space_before=4, space_after=2)

# ── HERO BLOCK ────────────────────────────────────────────────────────────────
p = doc.add_paragraph()
p.paragraph_format.space_before = Pt(6)
p.paragraph_format.space_after  = Pt(2)
r = p.add_run('PROPOSTA COMERCIAL')
r.bold = True; r.font.size = Pt(8); r.font.color.rgb = RED; r.font.all_caps = True
rPr = r._r.get_or_add_rPr()
sp = OxmlElement('w:spacing'); sp.set(qn('w:val'), '80'); rPr.append(sp)

p = doc.add_paragraph()
p.paragraph_format.space_before = Pt(2)
p.paragraph_format.space_after  = Pt(2)
r = p.add_run('Produção Musical, Mix & Master')
r.bold = True; r.font.size = Pt(22); r.font.color.rgb = DARK

p = doc.add_paragraph()
p.paragraph_format.space_before = Pt(0)
p.paragraph_format.space_after  = Pt(8)
r = p.add_run('com captação de vídeo em Chroma Key')
r.italic = True; r.font.size = Pt(13); r.font.color.rgb = GOLD

# Info chips via table
info_table = doc.add_table(rows=1, cols=4)
info_table.alignment = WD_TABLE_ALIGNMENT.LEFT
set_table_no_borders(info_table)
info_table.columns[0].width = Cm(4)
info_table.columns[1].width = Cm(4)
info_table.columns[2].width = Cm(4.25)
info_table.columns[3].width = Cm(4.25)

chips = [
    ('Serviço', '1 voz / 1 música'),
    ('Entrega', 'Estúdio Atoca Music'),
    ('Gravação', 'Voz solo profissional'),
    ('Vídeo', 'Chroma Key incluso'),
]
row = info_table.rows[0]
for i, (label, val) in enumerate(chips):
    cell = row.cells[i]
    set_cell_bg(cell, 247, 247, 247)
    no_borders(cell)
    set_cell_margin(cell, 80, 80, 110, 110)
    p_l = cell.paragraphs[0]
    p_l.paragraph_format.space_before = Pt(0)
    p_l.paragraph_format.space_after  = Pt(2)
    rl = p_l.add_run(label)
    rl.font.size = Pt(7.5); rl.font.color.rgb = LIGHT
    p_v = cell.add_paragraph()
    p_v.paragraph_format.space_before = Pt(0)
    p_v.paragraph_format.space_after  = Pt(0)
    rv = p_v.add_run(val)
    rv.bold = True; rv.font.size = Pt(9.5); rv.font.color.rgb = DARK

add_paragraph(doc, space_before=6, space_after=6)
add_hr(doc)

# ── SEÇÃO 01 APRESENTAÇÃO ─────────────────────────────────────────────────────
add_section_label(doc, '01 · Apresentação')
p = doc.add_paragraph()
p.paragraph_format.space_before = Pt(2)
p.paragraph_format.space_after  = Pt(6)
p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
add_run(p, 'A ', size=11, color=MID)
add_run(p, 'Atoca Music', bold=True, size=11, color=DARK)
add_run(p, ' é uma empresa especializada em produção musical, gravação, mixagem e masterização, com estúdio profissional em Brasília–DF. Apresentamos nossa proposta para a realização completa de '
           'produção musical de 1 (uma) voz, contemplando desde a gravação até a entrega do produto final mixado, masterizado e com captação de vídeo em fundo Chroma Key.', size=11, color=MID)

# Quote box
q_table = doc.add_table(rows=1, cols=1)
set_table_no_borders(q_table)
q_cell = q_table.cell(0,0)
set_cell_margin(q_cell, 100, 100, 160, 120)
tc = q_cell._tc
tcPr = tc.get_or_add_tcPr()
tcBorders = OxmlElement('w:tcBorders')
left_b = OxmlElement('w:left')
left_b.set(qn('w:val'), 'single')
left_b.set(qn('w:sz'), '16')
left_b.set(qn('w:color'), 'C27B00')
tcBorders.append(left_b)
tcPr.append(tcBorders)
set_cell_bg(q_cell, 254, 250, 240)
qp = q_cell.paragraphs[0]
qp.paragraph_format.space_before = Pt(2)
qp.paragraph_format.space_after  = Pt(2)
qr = qp.add_run('"Do estúdio para a tela — voz, som e imagem com a qualidade que a sua música merece."')
qr.italic = True; qr.font.size = Pt(11); qr.font.color.rgb = MID

add_paragraph(doc, space_before=6, space_after=6)
add_hr(doc)

# ── SEÇÃO 02 ESCOPO DOS SERVIÇOS ──────────────────────────────────────────────
add_section_label(doc, '02 · Escopo dos Serviços Incluídos')

cards = [
    ('🎤', 'Gravação de Voz',
     'Captação da voz solo em estúdio acústico tratado, com microfone condensador de alta qualidade, pré-amplificador profissional e monitoramento em tempo real.'),
    ('🎚️', 'Produção Musical',
     'Produção da base musical completa (arranjo, instrumentação e direção artística), alinhada ao estilo e referência do artista.'),
    ('🎛️', 'Mixagem (Mix)',
     'Tratamento individual de todas as trilhas, equalização, compressão, efeitos e balanceamento final para garantir clareza, profundidade e impacto sonoro.'),
    ('💿', 'Masterização (Master)',
     'Finalização do produto para distribuição digital e física, com loudness adequado às plataformas de streaming (Spotify, Deezer, YouTube, etc.).'),
    ('🎬', 'Captação em Chroma Key',
     'Gravação de vídeo em fundo verde (Chroma Key) com iluminação profissional, câmera HD/4K, enquadramentos artísticos e entrega do arquivo bruto editado.'),
    ('📦', 'Entrega dos Arquivos',
     'Entrega do áudio master em .WAV (alta resolução) e .MP3 (320kbps), mais o vídeo em formato digital, em até [prazo acordado] dias úteis após a sessão de gravação.'),
]

# 2 cards per row
for row_i in range(0, len(cards), 2):
    t = doc.add_table(rows=1, cols=2)
    set_table_no_borders(t)
    t.alignment = WD_TABLE_ALIGNMENT.LEFT
    t.columns[0].width = Cm(8.25)
    t.columns[1].width = Cm(8.25)
    for col_i in range(2):
        idx = row_i + col_i
        if idx >= len(cards): break
        icon, title, desc = cards[idx]
        cell = t.rows[0].cells[col_i]
        set_cell_bg(cell, 247, 247, 247)
        no_borders(cell)
        set_cell_margin(cell, 100, 100, 140, 140)
        # icon+title
        p_t = cell.paragraphs[0]
        p_t.paragraph_format.space_before = Pt(0)
        p_t.paragraph_format.space_after  = Pt(4)
        add_run(p_t, f'{icon}  {title}', bold=True, size=10.5, color=RED)
        # desc
        p_d = cell.add_paragraph(desc)
        p_d.paragraph_format.space_before = Pt(0)
        p_d.paragraph_format.space_after  = Pt(0)
        for run in p_d.runs:
            run.font.size = Pt(9.5); run.font.color.rgb = MID
    add_paragraph(doc, space_before=3, space_after=3)

add_paragraph(doc, space_before=4, space_after=4)
add_hr(doc)

# ── SEÇÃO 03 TABELA DE PREÇOS ─────────────────────────────────────────────────
add_section_label(doc, '03 · Tabela de Preços — por Música')

price_table = doc.add_table(rows=1, cols=3)
set_table_no_borders(price_table)
price_table.alignment = WD_TABLE_ALIGNMENT.LEFT
price_table.columns[0].width = Cm(8.5)
price_table.columns[1].width = Cm(4.5)
price_table.columns[2].width = Cm(3.5)

# Header row
hrow = price_table.rows[0]
hdata = [('SERVIÇO', WD_ALIGN_PARAGRAPH.LEFT), ('DESCRIÇÃO', WD_ALIGN_PARAGRAPH.CENTER), ('VALOR', WD_ALIGN_PARAGRAPH.RIGHT)]
for ci, (htxt, halign) in enumerate(hdata):
    cell = hrow.cells[ci]
    set_cell_bg(cell, 26, 26, 26)
    no_borders(cell)
    set_cell_margin(cell, 90, 90, 120, 120)
    ph = cell.paragraphs[0]
    ph.alignment = halign
    ph.paragraph_format.space_before = Pt(0)
    ph.paragraph_format.space_after  = Pt(0)
    rh = ph.add_run(htxt)
    rh.bold = True; rh.font.size = Pt(9); rh.font.color.rgb = WHITE; rh.font.all_caps = True

items = [
    ('Produção Musical Completa', 'Arranjo, instrumentação e direção artística', 'R$ ______,00'),
    ('Gravação de Voz (1 voz)', 'Sessão em estúdio profissional', 'R$ ______,00'),
    ('Mixagem (Mix)', 'Tratamento e balanceamento das trilhas', 'R$ ______,00'),
    ('Masterização (Master)', 'Finalização para streaming e distribuição', 'R$ ______,00'),
    ('Captação Vídeo – Chroma Key', 'Gravação de vídeo em fundo verde', 'R$ 100,00'),
]

for i, (srv, desc, val) in enumerate(items):
    row = price_table.add_row()
    bg = (252, 252, 252) if i % 2 == 0 else (245, 245, 245)
    cells_data = [(srv, WD_ALIGN_PARAGRAPH.LEFT, DARK, True),
                  (desc, WD_ALIGN_PARAGRAPH.CENTER, MID, False),
                  (val, WD_ALIGN_PARAGRAPH.RIGHT, DARK, True)]
    for ci, (txt, align, col, bld) in enumerate(cells_data):
        cell = row.cells[ci]
        set_cell_bg(cell, *bg)
        no_borders(cell)
        set_cell_margin(cell, 90, 90, 120, 120)
        pp = cell.paragraphs[0]
        pp.alignment = align
        pp.paragraph_format.space_before = Pt(0)
        pp.paragraph_format.space_after  = Pt(0)
        rr = pp.add_run(txt)
        rr.bold = bld; rr.font.size = Pt(10); rr.font.color.rgb = col

# Separator row
sep_row = price_table.add_row()
for ci in range(3):
    cell = sep_row.cells[ci]
    set_cell_bg(cell, 192, 57, 43)
    no_borders(cell)
    set_cell_margin(cell, 4, 4, 120, 120)
    cell.paragraphs[0].paragraph_format.space_before = Pt(0)
    cell.paragraphs[0].paragraph_format.space_after  = Pt(0)

# Total row
total_row = price_table.add_row()
total_cells = [
    ('PACOTE COMPLETO (produção + vídeo)', WD_ALIGN_PARAGRAPH.LEFT),
    ('1 música · 1 voz · chroma key', WD_ALIGN_PARAGRAPH.CENTER),
    ('R$ ______,00', WD_ALIGN_PARAGRAPH.RIGHT),
]
for ci, (txt, align) in enumerate(total_cells):
    cell = total_row.cells[ci]
    set_cell_bg(cell, 26, 26, 26)
    no_borders(cell)
    set_cell_margin(cell, 100, 100, 120, 120)
    pp = cell.paragraphs[0]
    pp.alignment = align
    pp.paragraph_format.space_before = Pt(0)
    pp.paragraph_format.space_after  = Pt(0)
    rr = pp.add_run(txt)
    rr.bold = True
    rr.font.size = Pt(11 if ci == 2 else 10)
    rr.font.color.rgb = RGBColor(0xEF, 0x9F, 0x27) if ci == 2 else WHITE

add_paragraph(doc, space_before=3, space_after=2)
p_obs = doc.add_paragraph()
p_obs.paragraph_format.space_before = Pt(2)
p_obs.paragraph_format.space_after  = Pt(4)
add_run(p_obs, '* O valor de Captação em Chroma Key é fixo em R$ 100,00 por música. Os demais valores serão preenchidos conforme briefing e complexidade do projeto. '
               'Pagamento: 50% no agendamento da sessão e 50% na entrega dos arquivos finais.', size=8.5, color=LIGHT, italic=True)

add_hr(doc)

# ── SEÇÃO 04 PROCESSO DE TRABALHO ─────────────────────────────────────────────
add_section_label(doc, '04 · Como Funciona — Processo de Trabalho')

steps = [
    ('01', 'Briefing & Direção Artística',
     'Conversa inicial para entender o estilo, referências musicais e objetivo do artista. Definição do arranjo, tom e formato do vídeo.'),
    ('02', 'Produção da Base Musical',
     'Criação do arranjo e produção instrumental na DAW, com aprovação do artista antes da gravação.'),
    ('03', 'Sessão de Gravação de Voz',
     'Gravação da voz em estúdio tratado acusticamente. Direção vocal incluída. Múltiplas tomadas para garantir a melhor performance.'),
    ('04', 'Captação de Vídeo (Chroma Key)',
     'Gravação do vídeo em fundo verde com iluminação e câmera profissional. Pode ser realizado na mesma sessão da gravação de voz.'),
    ('05', 'Mix & Master',
     'Mixagem e masterização completas, com até 2 rodadas de revisão incluídas no pacote.'),
    ('06', 'Entrega dos Arquivos',
     'Entrega do áudio (.WAV e .MP3) e do vídeo em formato digital, prontos para distribuição e publicação.'),
]

for step_no, step_title, step_desc in steps:
    t = doc.add_table(rows=1, cols=2)
    set_table_no_borders(t)
    t.alignment = WD_TABLE_ALIGNMENT.LEFT
    t.columns[0].width = Cm(1.6)
    t.columns[1].width = Cm(14.9)

    num_cell = t.rows[0].cells[0]
    set_cell_bg(num_cell, 192, 57, 43)
    no_borders(num_cell)
    set_cell_margin(num_cell, 60, 60, 80, 80)
    pn = num_cell.paragraphs[0]
    pn.alignment = WD_ALIGN_PARAGRAPH.CENTER
    pn.paragraph_format.space_before = Pt(0)
    pn.paragraph_format.space_after  = Pt(0)
    rn = pn.add_run(step_no)
    rn.bold = True; rn.font.size = Pt(11); rn.font.color.rgb = WHITE

    desc_cell = t.rows[0].cells[1]
    set_cell_bg(desc_cell, 247, 247, 247)
    no_borders(desc_cell)
    set_cell_margin(desc_cell, 60, 60, 130, 120)
    pt = desc_cell.paragraphs[0]
    pt.paragraph_format.space_before = Pt(0)
    pt.paragraph_format.space_after  = Pt(2)
    add_run(pt, step_title + '  ', bold=True, size=10.5, color=DARK)
    pd = desc_cell.add_paragraph(step_desc)
    pd.paragraph_format.space_before = Pt(0)
    pd.paragraph_format.space_after  = Pt(0)
    for r in pd.runs: r.font.size = Pt(9.5); r.font.color.rgb = MID
    add_paragraph(doc, space_before=2, space_after=2)

add_paragraph(doc, space_before=4, space_after=4)
add_hr(doc)

# ── SEÇÃO 05 DIFERENCIAIS ─────────────────────────────────────────────────────
add_section_label(doc, '05 · Diferenciais Atoca Music')

diffs = [
    ('Estúdio profissional em Brasília',
     'Ambiente acusticamente tratado, com equipamentos de alta qualidade para captação de voz.'),
    ('Pacote completo em um só lugar',
     'Gravação, produção, mix, master e vídeo sem precisar contratar múltiplos fornecedores.'),
    ('Chroma Key por apenas R$ 100,00',
     'Valor acessível para que o artista tenha conteúdo visual profissional junto com o lançamento musical.'),
    ('Entrega para plataformas de streaming',
     'Master adequado aos padrões de Spotify, Deezer, YouTube Music, Apple Music e demais.'),
    ('Revisões incluídas',
     'Até 2 rodadas de ajuste no mix e master sem custo adicional.'),
    ('Experiência e credibilidade',
     'Equipe com experiência em eventos e produções de artistas com projeção local e nacional.'),
]

for i in range(0, len(diffs), 2):
    t = doc.add_table(rows=1, cols=2)
    set_table_no_borders(t)
    t.alignment = WD_TABLE_ALIGNMENT.LEFT
    t.columns[0].width = Cm(8.25)
    t.columns[1].width = Cm(8.25)
    for ci in range(2):
        idx = i + ci
        if idx >= len(diffs): break
        title, desc = diffs[idx]
        cell = t.rows[0].cells[ci]
        set_cell_bg(cell, 247, 247, 247)
        no_borders(cell)

        # left red border
        tc2 = cell._tc
        tcPr2 = tc2.get_or_add_tcPr()
        brd = OxmlElement('w:tcBorders')
        lb = OxmlElement('w:left')
        lb.set(qn('w:val'), 'single')
        lb.set(qn('w:sz'), '12')
        lb.set(qn('w:color'), 'C0392B')
        brd.append(lb)
        tcPr2.append(brd)

        set_cell_margin(cell, 80, 80, 140, 120)
        pt = cell.paragraphs[0]
        pt.paragraph_format.space_before = Pt(0)
        pt.paragraph_format.space_after  = Pt(3)
        add_run(pt, title, bold=True, size=10, color=DARK)
        pd = cell.add_paragraph(desc)
        pd.paragraph_format.space_before = Pt(0)
        pd.paragraph_format.space_after  = Pt(0)
        for r in pd.runs: r.font.size = Pt(9.5); r.font.color.rgb = MID
    add_paragraph(doc, space_before=3, space_after=3)

add_paragraph(doc, space_before=4, space_after=4)
add_red_hr(doc)

# ── RODAPÉ ────────────────────────────────────────────────────────────────────
add_paragraph(doc, space_before=6, space_after=4)

footer_t = doc.add_table(rows=1, cols=2)
set_table_no_borders(footer_t)
footer_t.alignment = WD_TABLE_ALIGNMENT.LEFT
footer_t.columns[0].width = Cm(11)
footer_t.columns[1].width = Cm(5.5)

left_cell = footer_t.rows[0].cells[0]
no_borders(left_cell)
set_cell_margin(left_cell, 40, 40, 0, 60)
p1 = left_cell.paragraphs[0]
p1.paragraph_format.space_before = Pt(0)
p1.paragraph_format.space_after  = Pt(2)
add_run(p1, 'ATOCA MUSIC', bold=True, size=12, color=DARK)
p2 = left_cell.add_paragraph()
p2.paragraph_format.space_before = Pt(0)
p2.paragraph_format.space_after  = Pt(2)
add_run(p2, 'CNPJ: 49.634.339/0001-39   ·   (61) 9828-7404', size=9, color=LIGHT)
p3 = left_cell.add_paragraph()
p3.paragraph_format.space_before = Pt(0)
p3.paragraph_format.space_after  = Pt(0)
add_run(p3, 'www.atocamusic.com.br', size=9, color=RED)
p4 = left_cell.add_paragraph()
p4.paragraph_format.space_before = Pt(0)
p4.paragraph_format.space_after  = Pt(0)
add_run(p4, 'Brasília – DF', size=9, color=LIGHT)

right_cell = footer_t.rows[0].cells[1]
no_borders(right_cell)
set_cell_margin(right_cell, 40, 40, 60, 0)

validity_t = doc.add_table(rows=0, cols=1)  # placeholder, do inline instead
# Just a validity badge paragraph
pv = right_cell.paragraphs[0]
pv.alignment = WD_ALIGN_PARAGRAPH.RIGHT
pv.paragraph_format.space_before = Pt(6)
pv.paragraph_format.space_after  = Pt(2)
add_run(pv, '✓  Proposta válida por 15 dias', bold=True, size=9.5, color=RGBColor(0x2E, 0x7D, 0x32))

pv2 = right_cell.add_paragraph()
pv2.alignment = WD_ALIGN_PARAGRAPH.RIGHT
pv2.paragraph_format.space_before = Pt(0)
pv2.paragraph_format.space_after  = Pt(0)
add_run(pv2, 'Brasília–DF, maio de 2026.', size=9, color=LIGHT, italic=True)

# ── SAVE ──────────────────────────────────────────────────────────────────────
path_docx = '/home/claude/Proposta_Producao_Musical_AtocaMusic.docx'
doc.save(path_docx)
print(f'Salvo: {path_docx}')
